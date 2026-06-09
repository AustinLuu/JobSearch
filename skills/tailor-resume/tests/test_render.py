"""Tests for render_docx helpers — contact-link parsing, cert de-dup, fit ladder."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))
import render_docx as r  # noqa: E402


def test_contact_token_email():
    disp, href = r._classify_contact_token("austinowenluu@gmail.com")
    assert href == "mailto:austinowenluu@gmail.com" and disp == "austinowenluu@gmail.com"


def test_contact_token_linkedin_github_labeled():
    disp, href = r._classify_contact_token("linkedin.com/in/austin-luu")
    assert disp == "LinkedIn" and href == "https://linkedin.com/in/austin-luu"
    disp, href = r._classify_contact_token("https://github.com/austinluu")
    assert disp == "GitHub" and href == "https://github.com/austinluu"


def test_contact_token_other_url_is_portfolio():
    disp, href = r._classify_contact_token("https://www.austinluu.me/")
    assert disp == "Portfolio"
    assert href == "https://www.austinluu.me/"        # href preserved as-is


def test_contact_token_label_pipe_override():
    assert r._classify_contact_token("Site|austinluu.me") == ("Site", "https://austinluu.me")


def test_contact_token_location_plain_and_phone_dropped():
    assert r._classify_contact_token("Toronto, Canada") == ("Toronto, Canada", None)
    assert r._classify_contact_token("+1-416-451-3338") == (None, None)   # phone dropped


def test_dedupe_certs_drops_education_overlap():
    edu = [{"degree": "Deep Learning Specialization",
            "institution": "DeepLearning.AI", "details": []}]
    certs = ["DeepLearning.AI Deep Learning Specialization — 2021–2022",
             "Project Management Professional (PMP), PMI — 2026",
             "SolidWorks CSWA Certified"]
    out = r._dedupe_certs(certs, edu)
    assert out == ["Project Management Professional (PMP), PMI — 2026",
                   "SolidWorks CSWA Certified"]


def test_dedupe_certs_keeps_distinct():
    edu = [{"degree": "B.Eng Mechatronics", "institution": "TMU", "details": []}]
    certs = ["AWS Solutions Architect", "PMP"]
    assert r._dedupe_certs(certs, edu) == ["AWS Solutions Architect", "PMP"]


def test_fit_ladder_is_loosest_to_tightest():
    # Contract the fit logic relies on: relaxed is loosest, compact the floor.
    assert r.FIT_LADDER == ["relaxed", "default", "snug", "compact"]
    assert r.COMPACT_DENSITY.body_size <= r.DEFAULT_DENSITY.body_size
    assert r.COMPACT_DENSITY.margin <= r.SNUG_DENSITY.margin <= r.DEFAULT_DENSITY.margin


def test_norm_strips_nonalnum():
    import render_template as rt
    assert rt._norm("FDA 510(k) / ISO 13485.") == "fda510kiso13485"


def test_char_scale_inserts_w_in_order():
    # w:w must sit before w:sz in CT_RPr or the docx fails schema validation.
    import render_template as rt
    from docx import Document
    from docx.oxml.ns import qn
    d = Document()
    p = d.add_paragraph()
    r = p.add_run("x"); r.font.size = __import__("docx").shared.Pt(11)
    rt._set_char_scale(r, 95)
    rpr = r._r.find(qn("w:rPr"))
    tags = [c.tag for c in rpr]
    assert qn("w:w") in tags
    assert tags.index(qn("w:w")) < tags.index(qn("w:sz"))


def test_split_trailing_year():
    import render_template as rt
    assert rt._split_trailing_year("Project Management Professional (PMP), PMI — 2026") == ("Project Management Professional (PMP), PMI", "2026")
    assert rt._split_trailing_year("SolidWorks CSWA Certified") == ("SolidWorks CSWA Certified", "")
    assert rt._split_trailing_year("AWS Certified Solutions Architect, 2024") == ("AWS Certified Solutions Architect", "2024")


def test_apply_budget_tier_trims_whole_items_only():
    import render_template as rt
    d = {"key_achievements": [1, 2, 3],
         "areas_of_expertise": ["a1", "a2", "a3", "a4", "a5", "a6", "a7", "a8", "a9"],
         "additional_skills": "Tool1, Tool2, Tool3",
         "skills": [{"category": "X", "items": ["i1", "i2"]}],
         "certifications": ["Cert A", "Cert B", "Cert C"],
         "experience": [{"title": "A", "company": "X", "context": "ctx-A", "bullets": ["b1", "b2", "b3", "b4"]},
                        {"title": "B", "company": "Y", "context": "ctx-B", "bullets": ["c1", "c2", "c3"]},
                        {"title": "C", "company": "Z", "context": "ctx-C", "bullets": ["d1"]}],
         "projects": [{"title": "P1"}, {"title": "P2"}]}
    tier = next(b for b in rt.BUDGET_LADDER if b.name == "one-project")  # ka2/exp3/top3/other1/proj1, drop_additional
    out = rt.apply_budget_tier(d, tier)
    assert len(out["key_achievements"]) == 2
    assert len(out["experience"]) == 3
    assert len(out["experience"][0]["bullets"]) == 3   # top role capped
    assert len(out["experience"][1]["bullets"]) == 1   # other role capped to floor
    assert len(out["projects"]) == 1
    assert out["experience"][0]["bullets"] == ["b1", "b2", "b3"]  # verbatim prefix, text never edited
    assert len(d["experience"][0]["bullets"]) == 4               # original not mutated
    # header levers: areas capped, context dropped, additional skills dropped (whole-item omission)
    assert out["areas_of_expertise"] == ["a1", "a2", "a3"]       # capped to areas_cap=3
    assert all(r["context"] == "" for r in out["experience"])    # context dropped at this tier
    assert out["additional_skills"] == "" and out["skills"] == []  # additional-skills line dropped
    assert out["certifications"] == ["Cert A"]                      # trailing certs dropped (cap 1), most-relevant kept
    assert d["areas_of_expertise"] == ["a1", "a2", "a3", "a4", "a5", "a6", "a7", "a8", "a9"]  # original intact
    assert d["additional_skills"] == "Tool1, Tool2, Tool3"       # original intact
    assert d["certifications"] == ["Cert A", "Cert B", "Cert C"]   # original intact


def test_full_tier_preserves_header():
    import render_template as rt
    d = {"key_achievements": [1, 2, 3],
         "areas_of_expertise": ["a1", "a2", "a3", "a4", "a5", "a6", "a7", "a8", "a9"],
         "experience": [{"title": "A", "company": "X", "context": "ctx-A", "bullets": ["b1"]}],
         "projects": [{"title": "P1"}, {"title": "P2"}]}
    full = next(b for b in rt.BUDGET_LADDER if b.name == "full")
    out = rt.apply_budget_tier(d, full)
    assert out["areas_of_expertise"] == d["areas_of_expertise"]   # full keeps all 9 areas
    assert out["experience"][0]["context"] == "ctx-A"             # full keeps context


def test_budget_ladder_loosest_to_tightest():
    import render_template as rt
    names = [b.name for b in rt.BUDGET_LADDER]
    assert names[0] == "rich-max" and names[-1] == "last-resort"
    projs = [b.max_projects for b in rt.BUDGET_LADDER]
    scales = [b.min_scale for b in rt.BUDGET_LADDER]
    areas = [b.areas_cap for b in rt.BUDGET_LADDER]
    certs = [b.max_certifications for b in rt.BUDGET_LADDER]
    assert projs[0] >= projs[-1]
    assert all(scales[i] >= scales[i + 1] for i in range(len(scales) - 1))   # condense harder as it tightens
    assert all(areas[i] >= areas[i + 1] for i in range(len(areas) - 1))      # header shrinks monotonically
    assert all(certs[i] >= certs[i + 1] for i in range(len(certs) - 1))      # certs shrink monotonically


def test_header_trimmed_before_content_floors():
    """The user-locked guarantee: trim header before dropping below the content
    floors (>=3 experiences, >=2 projects, top role >=3 bullets). So every tier
    that still honors all three floors must come before any tier that breaks one."""
    import render_template as rt
    def honors_floors(b):
        return b.max_experiences >= 3 and b.max_projects >= 2 and b.top_bullets >= 3
    flags = [honors_floors(b) for b in rt.BUDGET_LADDER]
    # once a floor is broken it stays broken: no floor-honoring tier appears after a broken one
    first_break = next((i for i, ok in enumerate(flags) if not ok), len(flags))
    assert all(flags[:first_break]), "a floor-breaking tier appears before a floor-honoring one"
    assert not any(flags[first_break:]), "floors should not be restored after being broken"
    # and the floor-honoring prefix must actually trim header along the way
    prefix = rt.BUDGET_LADDER[:first_break]
    assert prefix[0].areas_cap > prefix[-1].areas_cap, "header (areas) not trimmed before floors break"
    assert any(b.drop_context for b in prefix), "context never dropped before floors break"


def test_ladder_is_richest_first_monotonic():
    """Pack-the-page contract: the ladder is ordered richest -> leanest, so the
    cascade's first-fit is the FULLEST page. Content weight must be monotonically
    non-increasing in every capped dimension, or 'first that fits' could skip a
    fuller option."""
    import render_template as rt
    L = rt.BUDGET_LADDER
    for dim in ("max_experiences", "top_bullets", "other_bullets", "max_projects",
                "key_achievements", "areas_cap", "max_certifications", "min_scale"):
        vals = [getattr(b, dim) for b in L]
        assert all(vals[i] >= vals[i + 1] for i in range(len(vals) - 1)), \
            f"{dim} is not monotonically non-increasing down the ladder: {vals}"


def test_growth_tiers_exceed_floors():
    """The richest tiers must actually GROW past the floors (else there's nothing
    to fill white space with). At least one tier pulls >3 experiences, >2 projects,
    and >3 top bullets — the material the cascade grows into."""
    import render_template as rt
    L = rt.BUDGET_LADDER
    assert max(b.max_experiences for b in L) >= 5, "no tier grows experiences past the floor"
    assert max(b.max_projects for b in L) >= 3, "no tier grows projects past the floor"
    assert max(b.top_bullets for b in L) >= 5, "no tier grows top-role bullets past the floor"
    # the richest tier keeps the full header (nothing trimmed yet)
    rm = L[0]
    assert not rm.drop_context and not rm.drop_additional and rm.areas_cap >= 9


def test_apply_budget_tier_growth_is_graceful_noop_when_cv_is_small():
    """A tier asking for more items than the CV supplies just yields what's there
    — growth never fabricates entries."""
    import render_template as rt
    d = {"experience": [{"title": "A", "bullets": ["b1", "b2"]},
                        {"title": "B", "bullets": ["b1"]}],
         "projects": [{"title": "P1", "bullets": ["x"]}],
         "key_achievements": [{"label": "k", "text": "t"}]}
    rich = next(b for b in rt.BUDGET_LADDER if b.name == "rich-max")
    out = rt.apply_budget_tier(d, rich)
    assert len(out["experience"]) == 2          # only 2 exist; cap of 5 is a no-op
    assert len(out["projects"]) == 1            # only 1 exists; cap of 4 is a no-op
    assert out["experience"][0]["bullets"] == ["b1", "b2"]  # fewer than top cap: unchanged


if __name__ == "__main__":
    fns = [v for k, v in sorted(globals().items()) if k.startswith("test_") and callable(v)]
    for fn in fns:
        fn()
        print(f"PASS {fn.__name__}")
    print(f"\n{len(fns)} tests passed")
