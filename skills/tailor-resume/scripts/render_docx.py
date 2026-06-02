"""
render_docx.py — render a TailoredResume (Filter 3 output) to an ATS-safe .docx.

ATS rules baked in (build-plan Phase 3, step 5):
  * US Letter, 1" margins, Arial.
  * The contact line lives in the document BODY, never a header/footer region
    (ATS parsers routinely drop header/footer text).
  * No columns, no text boxes, no tables-used-for-layout. A pure linear flow of
    headings, paragraphs, and real bullet-styled list items parses cleanly.
  * Real heading + bullet styles, never manual unicode bullets.

Template handling:
  * If a base template exists (default ~/Documents/JobSearch/templates/template_0.docx,
    or --template PATH), it is OPENED and its body cleared, so the template's
    own styles/fonts/margins are preserved and content is filled into them —
    "load and populate, don't re-theme per job."
  * If no template exists, a fresh document is built from the canonical style
    spec below (identical to what make_template.py bakes into template_0.docx),
    so the pipeline still produces consistent output out of the box.

Usage:
    python render_docx.py --tailored tailored.json --out resume.docx [--template t.docx]
"""

from __future__ import annotations

import argparse
import json
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

try:
    from . import common
except ImportError:  # pragma: no cover
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    import common  # type: ignore


# Canonical style spec — the single source of truth shared with make_template.py.
BODY_FONT = "Arial"
BODY_SIZE = 10.5          # pt
NAME_SIZE = 18            # pt
SECTION_SIZE = 12         # pt
LETTER_W, LETTER_H = 12240, 15840   # DXA (8.5 x 11")
MARGIN = 1440                        # DXA (1")
RULE_COLOR = "2E75B6"                # section-heading underline
INK = RGBColor(0x00, 0x00, 0x00)


# --------------------------------------------------------------------------
# Density profiles. "one page" is achieved first by CONTENT selection in the
# filters (SKILL.md) and only fine-tuned here. COMPACT trades whitespace for
# room but stays above a readability/ATS floor: never below ~10pt body or
# ~0.6" margins. Editing the module constants above changes the DEFAULT
# profile (the values are seeded from them), preserving the README's
# "edit BODY_SIZE / MARGIN" guidance.
# --------------------------------------------------------------------------

@dataclass(frozen=True)
class Density:
    margin: int = MARGIN          # DXA, all four sides
    body_size: float = BODY_SIZE  # pt
    name_size: float = NAME_SIZE  # pt
    section_size: float = SECTION_SIZE  # pt
    spacing: float = 1.0          # multiplier on every inter-paragraph gap


DEFAULT_DENSITY = Density()
# Airier than default — used to FILL a page when content is light, so a short
# resume doesn't leave a big blank lower half.
RELAXED_DENSITY = Density(name_size=20, section_size=12.5, spacing=1.3)
# Intermediate: trims margins/whitespace before reaching the readability floor.
SNUG_DENSITY = Density(margin=1080, name_size=17, section_size=11.5, spacing=0.75)
# Readability/ATS floor — never tighten past this to force a page; flag instead.
COMPACT_DENSITY = Density(margin=900, body_size=10.0, name_size=15,
                          section_size=11, spacing=0.5)
DENSITIES = {"relaxed": RELAXED_DENSITY, "default": DEFAULT_DENSITY,
             "snug": SNUG_DENSITY, "compact": COMPACT_DENSITY}
# Loosest -> tightest. render_fit picks the LOOSEST profile that fits max_pages,
# so the page is filled as much as possible (minimizes bottom whitespace) without
# ever dropping below COMPACT.
FIT_LADDER = ["relaxed", "default", "snug", "compact"]

# Active profile the helpers read; render() sets it per call.
_S = DEFAULT_DENSITY


def _sp(points: float):
    """Inter-paragraph spacing, scaled by the active density profile."""
    return Pt(points * _S.spacing)


# --------------------------------------------------------------------------
# Low-level docx helpers
# --------------------------------------------------------------------------

def _clear_body(doc: Document) -> None:
    """Remove all paragraphs/tables but keep the trailing sectPr (page setup)."""
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _set_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = common_dxa(LETTER_W)
        section.page_height = common_dxa(LETTER_H)
        section.top_margin = common_dxa(_S.margin)
        section.bottom_margin = common_dxa(_S.margin)
        section.left_margin = common_dxa(_S.margin)
        section.right_margin = common_dxa(_S.margin)


def common_dxa(dxa: int):
    from docx.shared import Twips
    return Twips(dxa)


def _ensure_default_font(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(_S.body_size)
    # Force the East-Asian slot too so the font sticks across renderers.
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), BODY_FONT)


# --- Hyperlinks + contact-line parsing -------------------------------------
from docx.opc.constants import RELATIONSHIP_TYPE as _RT  # noqa: E402
import re as _re  # noqa: E402

_EMAIL_RE = _re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
_DOMAIN_RE = _re.compile(r"^(https?://)?(www\.)?[a-z0-9.-]+\.[a-z]{2,}(/\S*)?$", _re.I)
_PHONE_RE = _re.compile(r"^[+(]?\d[\d\s().+\-]{6,}$")


def _classify_contact_token(tok: str):
    """One contact token -> (display_text, href_or_None), or (None, None) to DROP.

    'Label|URL' -> hyperlink shown as Label; phone -> dropped; email -> mailto with
    the address shown; LinkedIn/GitHub URL -> shown as that word; any other URL ->
    'Portfolio'; anything else (location) -> plain text.
    """
    t = tok.strip()
    if "|" in t:
        label, _, url = t.partition("|")
        label, url = label.strip(), url.strip()
        if label and url:
            href = url if url.lower().startswith(("http://", "https://")) else "https://" + url
            return label, href
    if "@" not in t and _PHONE_RE.match(t):
        return None, None
    if _EMAIL_RE.match(t):
        return t, "mailto:" + t
    if _DOMAIN_RE.match(t):
        href = t if t.lower().startswith(("http://", "https://")) else "https://" + t
        low = t.lower()
        if "linkedin." in low:
            return "LinkedIn", href
        if "github." in low:
            return "GitHub", href
        return "Portfolio", href
    return t, None


def _add_hyperlink(paragraph, href: str, text: str, *, size=None):
    """Append a real w:hyperlink run (blue, underlined) to a paragraph."""
    r_id = paragraph.part.relate_to(href, _RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), BODY_FONT)
    rpr.append(rf)
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    if size:
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(round(size * 2)))); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return link


# --- Certification / education de-duplication ------------------------------
_CERT_STOP = {
    "the", "a", "an", "of", "and", "for", "in", "on", "at", "to", "with", "by",
    "specialization", "certified", "certificate", "certification", "professional",
}


def _content_words(s) -> set:
    return {
        w for w in _re.findall(r"[a-z0-9+#]+", str(s).lower())
        if w not in _CERT_STOP and not w.isdigit() and len(w) > 1
    }


def _dedupe_certs(certs, education):
    """Drop any certification already represented in Education (its distinctive
    words are a subset of an education entry) — e.g. a DL Specialization that
    cv.md lists under BOTH Education and Certifications. Keeps genuinely distinct
    certs (PMP, CSWA, ...)."""
    edu_words = set()
    for e in education or []:
        edu_words |= _content_words(f"{e.get('degree','')} {e.get('institution','')}")
        for d in e.get("details", []):
            edu_words |= _content_words(d)
    out = []
    for c in certs or []:
        c = str(c).strip()
        if not c:
            continue
        cw = _content_words(c)
        if cw and cw <= edu_words:
            continue
        out.append(c)
    return out


# Per the CT_PPr schema, w:pBdr must appear before any of these siblings.
_PPR_AFTER_PBDR = (
    "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap",
    "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr",
    "w:pPrChange",
)


def _fix_settings(doc: Document) -> None:
    """python-docx's default template ships <w:zoom w:val="bestFit"/> with no
    w:percent, which fails strict OOXML validation. Normalize it."""
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None:
        # w:percent is required; w:val="bestFit" without it is invalid.
        if zoom.get(qn("w:percent")) is None:
            zoom.set(qn("w:percent"), "100")
        if zoom.get(qn("w:val")) is not None:
            del zoom.attrib[qn("w:val")]


def _add_bottom_border(paragraph) -> None:
    """Paragraph bottom border used as a section rule (NOT a table divider)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), RULE_COLOR)
    borders.append(bottom)
    # pBdr has a fixed position in the CT_PPr child order; insert it ahead of
    # any element that the schema requires to follow it (spacing/ind/jc/rPr/...).
    p_pr.insert_element_before(borders, *_PPR_AFTER_PBDR)


def _para(doc, text="", *, size=None, bold=False, italic=False,
          align=None, space_before=0, space_after=2, color=INK):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = _sp(space_before)
    pf.space_after = _sp(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = BODY_FONT
        run.font.size = Pt(_S.body_size if size is None else size)
        run.font.color.rgb = color
    return p


def _section_heading(doc, text):
    p = _para(doc, text.upper(), size=_S.section_size, bold=True,
              space_before=10, space_after=3)
    _add_bottom_border(p)
    return p


def _bullet(doc, text):
    # Built-in 'List Bullet' style is the ATS-safe way (real list, no glyph hack).
    try:
        p = doc.add_paragraph(style="List Bullet")
    except KeyError:
        p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(_S.body_size)
    p.paragraph_format.space_after = _sp(2)
    return p


# --------------------------------------------------------------------------
# Content rendering
# --------------------------------------------------------------------------

def _render_role(doc, role: dict):
    title = role.get("title", "")
    company = role.get("company", "")
    head = f"{title} — {company}".strip(" —")
    p = _para(doc, head, bold=True, space_before=6, space_after=0)
    meta_bits = [b for b in (role.get("location"), role.get("dates")) if b]
    if meta_bits:
        _para(doc, " · ".join(meta_bits), italic=True, size=_S.body_size - 0.5, space_after=1)
    if role.get("context"):
        _para(doc, role["context"], italic=True, space_after=2)
    for b in role.get("bullets", []):
        if b and b.strip():
            _bullet(doc, b.strip())


def _render_contact(doc, contact: dict):
    """Centered name + a single contact row in the BODY (never a header/footer),
    with email and links hyperlinked, and a thin rule dividing the header."""
    _para(doc, contact.get("name", ""), size=_S.name_size, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = _sp(3)
    csize = _S.body_size - 0.5
    tokens = [t.strip() for t in str(contact.get("line", "")).split("·") if t.strip()]
    parsed = []
    for tok in tokens:
        disp, href = _classify_contact_token(tok)
        if disp is not None:
            parsed.append((disp, href))
    for i, (disp, href) in enumerate(parsed):
        if i:
            sep = cp.add_run("   ·   ")
            sep.font.name = BODY_FONT
            sep.font.size = Pt(csize)
            sep.font.color.rgb = INK
        if href:
            _add_hyperlink(cp, href, disp, size=csize)
        else:
            r = cp.add_run(disp)
            r.font.name = BODY_FONT
            r.font.size = Pt(csize)
            r.font.color.rgb = INK
    _add_bottom_border(cp)


def render(tailored: dict, out_path: Path, template_path: Path | None = None,
           density: str = "default") -> Path:
    global _S
    result = common.validate_tailored_resume(tailored)
    if not result.ok:
        raise ValueError("TailoredResume failed validation:\n  - " + "\n  - ".join(result.errors))

    _S = DENSITIES.get(density, DEFAULT_DENSITY)
    try:
        return _render_body(tailored, out_path, template_path)
    finally:
        _S = DEFAULT_DENSITY  # never leak a profile to a later call in-process


def _render_body(tailored: dict, out_path: Path, template_path: Path | None) -> Path:
    if template_path and Path(template_path).exists():
        doc = Document(str(template_path))
        _clear_body(doc)
    else:
        doc = Document()
    _ensure_default_font(doc)
    _set_page(doc)
    _fix_settings(doc)
    contact = tailored["contact"]
    # Centered name + single hyperlinked contact row, both in the BODY.
    _render_contact(doc, contact)

    if tailored.get("summary"):
        _section_heading(doc, "Summary")
        _para(doc, tailored["summary"], space_after=4)

    skills = tailored.get("skills", [])
    if skills:
        _section_heading(doc, "Skills")
        for grp in skills:
            cat = grp.get("category", "")
            items = ", ".join(grp.get("items", []))
            p = doc.add_paragraph()
            p.paragraph_format.space_after = _sp(1)
            if cat:
                r = p.add_run(f"{cat}: ")
                r.bold = True
                r.font.name = BODY_FONT
                r.font.size = Pt(_S.body_size)
            r2 = p.add_run(items)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(_S.body_size)

    if tailored.get("experience"):
        _section_heading(doc, "Experience")
        for role in tailored["experience"]:
            _render_role(doc, role)

    if tailored.get("projects"):
        _section_heading(doc, "Projects")
        for role in tailored["projects"]:
            _render_role(doc, role)

    if tailored.get("education"):
        _section_heading(doc, "Education")
        for edu in common.strip_undergrad_achievements(tailored["education"]):
            head = ", ".join([b for b in (edu.get("degree"), edu.get("institution")) if b])
            _para(doc, head, bold=True, space_before=4, space_after=0)
            meta = [b for b in (edu.get("location"), edu.get("dates")) if b]
            if meta:
                _para(doc, " · ".join(meta), italic=True, size=_S.body_size - 0.5, space_after=1)
            for d in edu.get("details", []):
                if d and d.strip():
                    _bullet(doc, d.strip())

    certs = _dedupe_certs(tailored.get("certifications", []), tailored.get("education", []))
    if certs:
        _section_heading(doc, "Certifications / Additional")
        for c in certs:
            _bullet(doc, c)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# --------------------------------------------------------------------------
# Page measurement + one-page fit. Pagination is measured with LibreOffice
# (the same engine the docx validator uses); it is a close PROXY for Word's
# pagination, not identical, so fit targets a comfortable page, not a brim-full
# one. Requires `soffice` plus either `pypdf` or `pdfinfo`; if neither is
# present the fit step degrades gracefully (renders default, reports unknown).
# --------------------------------------------------------------------------

def count_pages(docx_path: Path | str) -> int | None:
    """Return the rendered page count, or None if it can't be measured."""
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as td:
        try:
            subprocess.run(
                ["soffice", "--headless",
                 f"-env:UserInstallation=file://{td}/profile",
                 "--convert-to", "pdf", "--outdir", td, str(docx_path)],
                check=True, capture_output=True, timeout=180,
            )
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            return None
        pdf = Path(td) / (docx_path.stem + ".pdf")
        if not pdf.exists():
            return None
        try:
            from pypdf import PdfReader
            return len(PdfReader(str(pdf)).pages)
        except Exception:
            pass
        try:
            out = subprocess.run(["pdfinfo", str(pdf)], capture_output=True,
                                 text=True, timeout=30).stdout
            for line in out.splitlines():
                if line.startswith("Pages:"):
                    return int(line.split()[1])
        except Exception:
            return None
    return None


def render_fit(tailored: dict, out_path: Path, template_path: Path | None = None,
               max_pages: int = 1) -> dict:
    """Render at the LOOSEST density that fits `max_pages`, walking the FIT_LADDER
    (relaxed -> default -> snug -> compact). Picking the loosest fit fills the page
    as much as possible, so a trimmed resume doesn't leave a big blank lower half.
    NEVER tightens past COMPACT; if even compact overflows, it leaves the compact
    best-effort on disk and reports fit_ok=False so the skill trims CONTENT rather
    than shrinking type further."""
    out_path = Path(out_path)
    attempts: list[tuple[str, int | None]] = []
    for name in FIT_LADDER:
        render(tailored, out_path, template_path, density=name)
        pages = count_pages(out_path)
        if pages is None:
            # Measurement unavailable; settle on default density, report unknown.
            render(tailored, out_path, template_path, density="default")
            return {"pages": None, "density": "default", "fit_ok": None,
                    "max_pages": max_pages,
                    "note": "page count unmeasurable (need soffice + pypdf/pdfinfo); "
                            "rendered at default density without verifying length"}
        attempts.append((name, pages))
        if pages <= max_pages:
            return {"pages": pages, "density": name, "fit_ok": True,
                    "max_pages": max_pages, "ladder": attempts}
    # Nothing fit even at the tightest rung (compact, now on disk).
    name, pages = attempts[-1]
    return {
        "pages": pages, "density": name, "fit_ok": False, "max_pages": max_pages,
        "ladder": attempts,
        "note": (f"still {pages} page(s) at the compact readability floor. Do NOT shrink "
                 f"further — return to Filter 3 and cut the least-relevant WHOLE roles/"
                 f"bullets (honest omission), or accept {pages} pages and pass "
                 f"--max-pages {pages}, marking the job 'needs human check'."),
    }


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Render a TailoredResume to ATS-safe .docx")
    ap.add_argument("--tailored", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--template", default=None,
                    help="base template .docx (default: locked templates/template_0.docx if present)")
    ap.add_argument("--max-pages", type=int, default=None,
                    help="fit the resume to at most this many pages (default 1 when set). "
                         "Omit to render once at --density with no length check.")
    ap.add_argument("--density", choices=("relaxed", "default", "snug", "compact"), default="default",
                    help="layout density when NOT using --max-pages")
    args = ap.parse_args(argv)

    tailored = common.load_json(args.tailored)
    tpl = Path(args.template) if args.template else common.TEMPLATE_PATH
    template = tpl if tpl.exists() else None

    if args.max_pages is not None:
        status = render_fit(tailored, Path(args.out), template, max_pages=args.max_pages)
        print(json.dumps({"out": str(args.out), **status}))
    else:
        out = render(tailored, Path(args.out), template, density=args.density)
        print(f"wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
