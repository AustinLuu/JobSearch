"""
render_template.py — fill the user's designed template_0.docx in place.

Unlike render_docx.py (which builds an ATS resume from scratch in a fixed style),
this renderer OPENS template_0.docx and swaps the candidate's content into its
existing structure, so the template's embedded fonts, section rules, and 3-column
skills layout are preserved exactly ("load and populate, don't re-theme").

It does this by:
  * editing text in place for the single-slot lines (name, role title, contact,
    summary, additional skills),
  * rebuilding the tab-aligned lines (company/location, title/dates, project,
    university) with a real RIGHT tab stop, reusing the template's own run
    formatting (so the embedded serif font carries over), and
  * cloning the template's repeatable blocks (skill bullet, achievement line,
    role block, project block, education entry) to fit variable-length content.

Template sections detected by their fixed heading text:
  Name / Role Title / Contact / Summary / AREAS OF EXPERTISE (3-col skills) /
  KEY ACHIEVEMENTS / PROFESSIONAL EXPERIENCE / PROJECTS / EDUCATION & CERTIFICATIONS
  / ADDITIONAL SKILLS.

Input: the TailoredResume dict (render_docx schema) plus a few template-specific
optional fields the filters supply:
  * role_title        — the headline under the name (per decision: the TARGET job title)
  * key_achievements  — [{ "label": "...", "text": "..." }, ...] (top 2-3)
  * areas_of_expertise— ["skill", ...] flat list for the 3-col block (else derived
                        from `skills`)
  * additional_skills — string for the ADDITIONAL SKILLS line (else derived)

ATS note: the template's skills are 3-column by the user's choice; `fill()` returns
an `ats_flags` list carrying that caveat so the skill can record it per run.
"""

from __future__ import annotations

import argparse
import copy
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt
from docx.text.paragraph import Paragraph

try:
    from . import common
except ImportError:  # pragma: no cover
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    import common  # type: ignore

try:
    from . import measure_fit
except ImportError:  # pragma: no cover
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    import measure_fit  # type: ignore

LINK_COLOR = "0563C1"

# One-line tightener tuning. WIDOW_MIN_SCALE is the hardest horizontal condense the
# widow-cleanup will apply. With the two-phase content packer now doing the real
# fitting, the tightener is only a GENTLE polish: a bullet within a few percent of
# fitting one line is nudged to ~93% (visually near-imperceptible, and uniform — all
# condensed lines land at the same width instead of spreading 82–92%); anything that
# would need more is left to WRAP naturally and the packer absorbs the extra line.
# SLOP is the extra condense (in percent) applied beyond the computed fit to absorb
# the LibreOffice-vs-Word measurement gap so a pulled-up line clears the wrap in Word
# instead of grazing it; with the higher floor it mainly guarantees condensed lines
# settle at the floor rather than above it.
WIDOW_MIN_SCALE = 0.93
SLOP = 7


# --------------------------------------------------------------------------
# Low-level element helpers (operate on python-docx Paragraph objects)
# --------------------------------------------------------------------------

def _text(p: Paragraph) -> str:
    return "".join(n.text for n in p._p.iter(qn("w:t")))


def _remove(p: Paragraph) -> None:
    p._p.getparent().remove(p._p)


def _clone_after(anchor: Paragraph, source: Paragraph) -> Paragraph:
    """Deep-copy `source` (with all formatting) and insert it right after `anchor`."""
    el = copy.deepcopy(source._p)
    anchor._p.addnext(el)
    return Paragraph(el, anchor._parent)


def _strip_content(p: Paragraph) -> None:
    """Remove all runs/hyperlinks (keep pPr: style, spacing, borders, numbering)."""
    for tag in ("w:r", "w:hyperlink"):
        for el in p._p.findall(qn(tag)):
            p._p.remove(el)


def _proto_rpr(run):
    """A deep-copied <w:rPr> from a template run, so new runs inherit its font."""
    rpr = run._r.find(qn("w:rPr"))
    return copy.deepcopy(rpr) if rpr is not None else None


def _set_bool(rpr, tag: str, on: bool):
    """Force bold/italic on/off inside a cloned rPr (so a bold proto can go regular)."""
    if rpr is None:
        return
    existing = rpr.find(qn(tag))
    if on and existing is None:
        rpr.append(OxmlElement(tag))
    elif not on and existing is not None:
        rpr.remove(existing)


def _add_run(p: Paragraph, text: str, proto_rpr, *, bold=None, italic=None):
    r = OxmlElement("w:r")
    if proto_rpr is not None:
        rpr = copy.deepcopy(proto_rpr)
        if bold is not None:
            _set_bool(rpr, "w:b", bold)
        if italic is not None:
            _set_bool(rpr, "w:i", italic)
        r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p._p.append(r)
    return r


def _add_tab(p: Paragraph, proto_rpr):
    r = OxmlElement("w:r")
    if proto_rpr is not None:
        r.append(copy.deepcopy(proto_rpr))
    r.append(OxmlElement("w:tab"))
    p._p.append(r)


def _add_hyperlink(p: Paragraph, href: str, text: str, proto_rpr):
    r_id = p.part.relate_to(
        href, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rpr = copy.deepcopy(proto_rpr) if proto_rpr is not None else OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), LINK_COLOR); rpr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t)
    link.append(r)
    p._p.append(link)


def _right_tab_stop(p: Paragraph, content_width_emu: int):
    # Clear any existing tab-stop defs, then add a single RIGHT stop at the margin.
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        for tabs in pPr.findall(qn("w:tabs")):
            pPr.remove(tabs)
    p.paragraph_format.tab_stops.add_tab_stop(Emu(content_width_emu), WD_TAB_ALIGNMENT.RIGHT)


def _set_split_line(p, left, right, protos, content_width, *,
                    left_bold=True, right_bold=False, right_italic=False,
                    right_href=None):
    """Rebuild a paragraph as 'LEFT <tab-to-right-margin> RIGHT', keeping pPr."""
    _strip_content(p)
    _right_tab_stop(p, content_width)
    _add_run(p, left, protos["bold"] if left_bold else protos["reg"],
             bold=left_bold, italic=False)
    _add_tab(p, protos["reg"])
    if right:
        if right_href:
            _add_hyperlink(p, right_href, right, protos["bold"] if right_bold else protos["reg"])
        else:
            _add_run(p, right, protos["bold"] if right_bold else protos["reg"],
                     bold=right_bold, italic=right_italic)


# --------------------------------------------------------------------------
# Contact-token classification (mirror of render_docx, kept local)
# --------------------------------------------------------------------------
import re as _re  # noqa: E402

_EMAIL_RE = _re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
_DOMAIN_RE = _re.compile(r"^(https?://)?(www\.)?[a-z0-9.-]+\.[a-z]{2,}(/\S*)?$", _re.I)
_PHONE_RE = _re.compile(r"^[+(]?\d[\d\s().+\-]{6,}$")


def _classify_contact_token(tok: str):
    """Map a contact token to (display, href), or (None, None) to DROP it.

    - 'Label|URL'  -> hyperlink shown as Label (explicit override).
    - phone number -> dropped (kept off the resume).
    - email        -> shown as the address, mailto-linked.
    - LinkedIn/GitHub URL -> shown as the word 'LinkedIn'/'GitHub', linked.
    - any other URL/domain -> shown as 'Portfolio', linked.
    - anything else (e.g. location) -> plain text.
    """
    t = tok.strip()
    if "|" in t:                                   # explicit Label|URL
        label, _, url = t.partition("|")
        label, url = label.strip(), url.strip()
        if label and url:
            href = url if url.lower().startswith(("http://", "https://")) else "https://" + url
            return label, href
    if "@" not in t and _PHONE_RE.match(t):        # drop phone numbers
        return None, None
    if _EMAIL_RE.match(t):
        return t, "mailto:" + t
    if _DOMAIN_RE.match(t):
        href = t if t.lower().startswith(("http://", "https://")) else "https://" + t
        low = t.lower()
        if "linkedin." in low:
            return "LinkedIn", href
        if "github." in low:
            return "GitHub", href
        return "Portfolio", href
    return t, None


# --------------------------------------------------------------------------
# Section location
# --------------------------------------------------------------------------

def _split_trailing_year(s: str):
    """Split a credential string into (name, year) when it ends with a 4-digit year
    (optionally after a dash/comma), else (s, ''). 'PMP, PMI — 2026' -> ('PMP, PMI','2026')."""
    m = _re.search(r"\s*[—–\-,]?\s*((?:19|20)\d{2})\s*$", s)
    if m:
        return s[:m.start()].rstrip(" ,—–-"), m.group(1)
    return s, ""


def _find(paras, predicate):
    for i, p in enumerate(paras):
        if predicate(_text(p).strip()):
            return i
    return None


def _derive_areas(tailored: dict, limit: int = 9):
    """Flat skills list for the 3-col block: first items across skill groups."""
    flat = []
    for grp in tailored.get("skills", []):
        for it in grp.get("items", []):
            if it not in flat:
                flat.append(it)
    return flat[:limit]


def _derive_additional(tailored: dict) -> str:
    """Remaining skill groups -> a single 'Additional Skills' line."""
    bits = []
    for grp in tailored.get("skills", []):
        cat = grp.get("category", "")
        items = ", ".join(grp.get("items", []))
        if items:
            bits.append(f"{cat}: {items}" if cat else items)
    return " | ".join(bits)


# --------------------------------------------------------------------------
# Fill
# --------------------------------------------------------------------------

def fill(tailored: dict, out_path: Path, template_path: Path) -> dict:
    res = common.validate_tailored_resume(tailored)
    if not res.ok:
        raise ValueError("TailoredResume failed validation:\n  - " + "\n  - ".join(res.errors))

    doc = Document(str(template_path))
    sec = doc.sections[-1]
    content_width = sec.page_width - sec.left_margin - sec.right_margin
    P = doc.paragraphs

    # --- prototype run formatting (captured from the live template) -----------
    i_company = _find(P, lambda t: t.startswith("COMPANY NAME"))
    i_module = _find(P, lambda t: t.startswith("Module studied"))
    i_summary = _find(P, lambda t: t.startswith("Biography Summary"))
    protos = {
        "bold": _proto_rpr(P[i_company].runs[0]),
        "reg": _proto_rpr(P[i_module].runs[0]),
        "ital": _proto_rpr(P[i_summary].runs[0]),
    }

    # --- single-slot lines ----------------------------------------------------
    contact = tailored["contact"]
    P[0].runs[0].text = contact.get("name", "")                       # Name
    # Role title intentionally NOT rendered (template update: no headline line under
    # the name). Strip paragraph 1 to clear any placeholder and leave it as the empty
    # spacer the template keeps between the name and the contact line. A `role_title`
    # in the tailored JSON is accepted but ignored.
    _strip_content(P[1])

    i_contact = _find(P, lambda t: "•" in t and "@" not in t and t[:4].isalpha()) or 2
    cp = P[i_contact]
    _strip_content(cp)
    tokens = [t.strip() for t in str(contact.get("line", "")).split("·") if t.strip()]
    parsed = []
    for tok in tokens:
        disp, href = _classify_contact_token(tok)
        if disp is not None:
            parsed.append((disp, href))
    for j, (disp, href) in enumerate(parsed):
        if j:
            _add_run(cp, " • ", protos["reg"])
        if href:
            _add_hyperlink(cp, href, disp, protos["reg"])
        else:
            _add_run(cp, disp, protos["reg"])

    sp = P[i_summary]
    _strip_content(sp)
    _add_run(sp, tailored.get("summary", ""), protos["ital"], italic=True)

    # --- AREAS OF EXPERTISE: 3-col skills (preserve the column sectPr) ---------
    i_ae = _find(P, lambda t: t.startswith("AREAS OF EXPERTISE"))
    i_ka = _find(P, lambda t: t.startswith("KEY ACHIEVEMENTS"))
    skill_paras = [p for p in P[i_ae + 1:i_ka] if p._p.find(qn("w:pPr")) is not None
                   and p._p.find(qn("w:pPr")).find(qn("w:numPr")) is not None]
    col_sectpr = None
    for p in skill_paras:
        spr = p._p.find(qn("w:pPr")).find(qn("w:sectPr"))
        if spr is not None:
            col_sectpr = copy.deepcopy(spr)        # the 3-column section break
    bullet_proto = copy.deepcopy(skill_paras[0]._p)
    anchor = skill_paras[0]
    prev = anchor.getprevious() if hasattr(anchor, "getprevious") else None
    insert_after = Paragraph(skill_paras[0]._p.getprevious(), doc)   # the SECTBREAK empty para
    for p in skill_paras:
        _remove(p)
    areas = tailored.get("areas_of_expertise") or _derive_areas(tailored)
    new_skill_paras = []
    cur = insert_after
    for skill in areas:
        el = copy.deepcopy(bullet_proto)
        # strip any sectPr the clone source carried; we re-add to the LAST one only
        pPr = el.find(qn("w:pPr"))
        spr = pPr.find(qn("w:sectPr")) if pPr is not None else None
        if spr is not None:
            pPr.remove(spr)
        cur._p.addnext(el)
        np = Paragraph(el, doc)
        _strip_content(np)
        _add_run(np, skill, _proto_rpr_from_el(bullet_proto))
        new_skill_paras.append(np)
        cur = np
    if col_sectpr is not None and new_skill_paras:
        last_pPr = new_skill_paras[-1]._p.find(qn("w:pPr"))
        if last_pPr is None:
            last_pPr = new_skill_paras[-1]._p.makeelement(qn("w:pPr"), {})
            new_skill_paras[-1]._p.insert(0, last_pPr)
        last_pPr.append(col_sectpr)

    # --- KEY ACHIEVEMENTS ------------------------------------------------------
    P = doc.paragraphs
    i_ka = _find(P, lambda t: t.startswith("KEY ACHIEVEMENTS"))
    i_pe = _find(P, lambda t: t.startswith("PROFESSIONAL EXPERIENCE"))
    ach_paras = [p for p in P[i_ka + 1:i_pe] if _text(p).strip()]
    ach_proto = copy.deepcopy(ach_paras[0]._p)
    anchor = Paragraph(ach_paras[0]._p.getprevious(), doc)
    for p in ach_paras:
        _remove(p)
    achievements = tailored.get("key_achievements") or []
    cur = anchor
    for a in achievements:
        el = copy.deepcopy(ach_proto)
        cur._p.addnext(el)
        np = Paragraph(el, doc)
        _strip_content(np)
        label = a.get("label", "").strip()
        if label:
            _add_run(np, f"{label}: ", protos["bold"], bold=True)
        _add_run(np, a.get("text", ""), protos["reg"], bold=False)
        cur = np

    # --- PROFESSIONAL EXPERIENCE ----------------------------------------------
    P = doc.paragraphs
    i_pe = _find(P, lambda t: t.startswith("PROFESSIONAL EXPERIENCE"))
    i_pr = _find(P, lambda t: t.startswith("PROJECT"))
    block = P[i_pe + 1:i_pr]
    comp_proto = copy.deepcopy(_first(block, lambda t: t.startswith("COMPANY NAME"))._p)
    title_proto = copy.deepcopy(_after(block, "COMPANY NAME")._p)
    bullet_proto_e = copy.deepcopy(_first_bullet(block)._p)
    gap_proto = copy.deepcopy(_first_empty(block)._p)
    anchor = Paragraph(block[0]._p.getprevious(), doc)
    for p in block:
        _remove(p)
    cur = anchor
    for role in tailored.get("experience", []):
        cur = _emit_split(cur, comp_proto, role.get("company", ""),
                          role.get("location", ""), protos, content_width,
                          left_bold=True, right_bold=True)
        cur = _emit_split(cur, title_proto, role.get("title", ""),
                          role.get("dates", ""), protos, content_width,
                          left_bold=False, right_italic=True)
        for b in role.get("bullets", []):
            if b and b.strip():
                cur = _emit_bullet(cur, bullet_proto_e, b.strip(), protos["reg"])
        cur = _emit_clone(cur, gap_proto)

    # --- PROJECTS --------------------------------------------------------------
    P = doc.paragraphs
    i_pr = _find(P, lambda t: t.strip().rstrip(".") == "PROJECTS")
    i_ed = _find(P, lambda t: t.startswith("EDUCATION"))
    block = P[i_pr + 1:i_ed]
    pname_proto = copy.deepcopy(_first(block, lambda t: t.startswith("PROJECT NAME"))._p)
    pdate_proto = copy.deepcopy(_first(block, lambda t: t.startswith("Month year"))._p)
    pbullet_proto = copy.deepcopy(_first_bullet(block)._p)
    pgap_proto = copy.deepcopy(_first_empty(block)._p)
    anchor = Paragraph(block[0]._p.getprevious(), doc)
    for p in block:
        _remove(p)
    cur = anchor
    for proj in tailored.get("projects", []):
        cur = _emit_split(cur, pname_proto, proj.get("title", ""),
                          proj.get("location", "") or proj.get("dates", ""),
                          protos, content_width, left_bold=True, right_bold=True)
        if proj.get("location") and proj.get("dates"):
            cur = _emit_split(cur, pdate_proto, "", proj.get("dates", ""),
                              protos, content_width, right_italic=True)
        for b in proj.get("bullets", []):
            if b and b.strip():
                cur = _emit_bullet(cur, pbullet_proto, b.strip(), protos["reg"])
        cur = _emit_clone(cur, pgap_proto)

    # --- EDUCATION & CERTIFICATIONS -------------------------------------------
    from render_docx import _dedupe_certs  # reuse the same dedupe logic
    P = doc.paragraphs
    i_ed = _find(P, lambda t: t.startswith("EDUCATION"))
    i_as = _find(P, lambda t: t.startswith("ADDITIONAL SKILLS"))
    block = P[i_ed + 1:i_as]
    uni_proto = copy.deepcopy(_first(block, lambda t: t.startswith("UNIVERSITY"))._p)
    detail_proto = copy.deepcopy(_first(block, lambda t: t.startswith("Module"))._p)
    anchor = Paragraph(block[0]._p.getprevious(), doc)
    for p in block:
        _remove(p)
    cur = anchor
    for edu in common.strip_undergrad_achievements(tailored.get("education", [])):
        head = ", ".join([b for b in (edu.get("degree"), edu.get("institution")) if b])
        cur = _emit_split(cur, uni_proto, head, edu.get("dates", ""),
                          protos, content_width, left_bold=True, right_italic=True)
        for d in edu.get("details", []):
            if d and d.strip():
                cur = _emit_detail(cur, detail_proto, d.strip(), protos["reg"])
    for cert in _dedupe_certs(tailored.get("certifications", []), tailored.get("education", [])):
        # Render each cert as its OWN entry (bold name + right-aligned year if present),
        # matching the education-entry style — not the subordinate detail style, which
        # makes a cert look like a sub-line of the education entry above it.
        left, year = _split_trailing_year(cert)
        cur = _emit_split(cur, uni_proto, left, year, protos, content_width,
                          left_bold=True, right_italic=True)

    # --- ADDITIONAL SKILLS -----------------------------------------------------
    # Multi-line: one paragraph per skill category, the "Category:" label bolded.
    # Renders from the grouped `skills` data (category + items). Falls back to the
    # explicit `additional_skills` string as a single UNLABELED line only when no
    # groups exist. Whole-line omission only — categories and items are never
    # reworded, so integrity is untouched. drop_additional clears both `skills`
    # and `additional_skills` (apply_budget_tier), so lean tiers still drop the
    # whole block.
    P = doc.paragraphs
    i_as = _find(P, lambda t: t.startswith("ADDITIONAL SKILLS"))
    as_paras = [p for p in P[i_as + 1:] if _text(p).strip()]   # existing content line(s)

    groups = []
    for grp in (tailored.get("skills") or []):
        items = ", ".join(grp.get("items", []))
        if items:
            groups.append((grp.get("category", "").strip(), items))
    if not groups:
        add_text = tailored.get("additional_skills") or _derive_additional(tailored)
        if str(add_text).strip():
            groups = [("", str(add_text).strip())]   # single unlabeled line (fallback)

    if not groups:
        # nothing to show: remove the heading and its content line so no orphan
        # heading is left at the foot of the page (a fit lever, not a rewrite).
        for p in as_paras:
            _remove(p)
        _remove(P[i_as])
    else:
        # Clone the template's content paragraph once per category so every line
        # inherits the template's own style/spacing (don't re-theme).
        if as_paras:
            proto = copy.deepcopy(as_paras[0]._p)
            anchor = Paragraph(as_paras[0]._p.getprevious(), doc)
            for p in as_paras:
                _remove(p)
        else:
            proto = copy.deepcopy(P[i_as + 1]._p)
            anchor = P[i_as]
        cur = anchor
        for cat, items in groups:
            el = copy.deepcopy(proto)
            cur._p.addnext(el)
            np = Paragraph(el, doc)
            _strip_content(np)
            if cat:
                _add_run(np, f"{cat}: ", protos["bold"], bold=True)
            _add_run(np, items, protos["reg"], bold=False)
            cur = np

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return {"out": str(out_path),
            "ats_flags": ["Skills render in a 3-column section (template design, user-approved). "
                          "Some ATS parsers read multi-column text out of order — verify parsing "
                          "if applying through a strict ATS."]}


# --- small block helpers ---------------------------------------------------

def _proto_rpr_from_el(p_el):
    r = p_el.find(qn("w:r"))
    if r is None:
        return None
    rpr = r.find(qn("w:rPr"))
    return copy.deepcopy(rpr) if rpr is not None else None


def _first(block, pred):
    for p in block:
        if pred(_text(p).strip()):
            return p
    return None


def _after(block, prefix):
    for k, p in enumerate(block):
        if _text(p).strip().startswith(prefix):
            return block[k + 1]
    return None


def _first_bullet(block):
    for p in block:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            return p
    return None


def _first_empty(block):
    for p in block:
        if not _text(p).strip():
            return p
    return block[0]


def _emit_split(cur, proto_el, left, right, protos, content_width, **kw):
    el = copy.deepcopy(proto_el)
    cur._p.addnext(el)
    np = Paragraph(el, cur._parent)
    _set_split_line(np, left, right, protos, content_width, **kw)
    return np


def _emit_bullet(cur, proto_el, text, reg_rpr):
    el = copy.deepcopy(proto_el)
    cur._p.addnext(el)
    np = Paragraph(el, cur._parent)
    _strip_content(np)
    proto = _proto_rpr_from_el(proto_el)
    _add_run(np, text, proto if proto is not None else reg_rpr)
    return np


def _emit_detail(cur, proto_el, text, reg_rpr):
    el = copy.deepcopy(proto_el)
    cur._p.addnext(el)
    np = Paragraph(el, cur._parent)
    _strip_content(np)
    proto = _proto_rpr_from_el(proto_el)
    _add_run(np, text, proto if proto is not None else reg_rpr)
    return np


def _emit_clone(cur, proto_el):
    el = copy.deepcopy(proto_el)
    cur._p.addnext(el)
    return Paragraph(el, cur._parent)


def _set_char_scale(run, pct: int):
    """Apply horizontal glyph scaling (<w:w w:val=...>) to a run, in CT_RPr order."""
    rpr = run._r.get_or_add_rPr()
    existing = rpr.find(qn("w:w"))
    if existing is not None:
        rpr.remove(existing)
    w = OxmlElement("w:w")
    w.set(qn("w:val"), str(pct))
    # w:w sits after w:spacing and before w:kern/w:position/w:sz in CT_RPr.
    after = ("w:kern", "w:position", "w:sz", "w:szCs", "w:highlight", "w:u",
             "w:effect", "w:bdr", "w:shd", "w:fitText", "w:vertAlign", "w:rtl",
             "w:cs", "w:em", "w:lang", "w:eastAsianLayout", "w:specVanish")
    anchor = next((c for c in rpr if c.tag in {qn(t) for t in after}), None)
    if anchor is not None:
        anchor.addprevious(w)
    else:
        rpr.append(w)


def _norm(s: str) -> str:
    return _re.sub(r"[^a-z0-9]", "", s.lower())


def _pdf_lines(pdf_path):
    """Ordered [(text, x0, x1)] across pages. Words are clustered into visual lines
    by a vertical tolerance (so a bullet glyph sitting a couple points off its text
    still groups with it)."""
    import pdfplumber
    out = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            words = sorted(page.extract_words(use_text_flow=True), key=lambda w: (w["top"], w["x0"]))
            cur, cur_top = [], None
            for w in words:
                if cur_top is None or abs(w["top"] - cur_top) <= 7:   # ~half a line
                    cur.append(w)
                    cur_top = w["top"] if cur_top is None else cur_top
                else:
                    out.append(_line_tuple(cur))
                    cur, cur_top = [w], w["top"]
            if cur:
                out.append(_line_tuple(cur))
    return out


def _line_tuple(ws):
    ws = sorted(ws, key=lambda w: w["x0"])
    return (" ".join(w["text"] for w in ws),
            min(w["x0"] for w in ws), max(w["x1"] for w in ws))


def tighten_wrapped_lines(docx_path, *, min_scale: float = 0.90) -> dict:
    """Best-effort cosmetic pass: a TARGET paragraph (bullet, key-achievement line,
    education detail, additional-skills line, summary) that wraps to exactly TWO lines
    by a small amount gets a near-invisible horizontal condense (w:w) so it collapses to
    one line — ONLY when the needed scale stays >= the floor. Bigger overflows are left
    to wrap (a content-trim decision, not a squish-the-type one).

    Skips name/role/contact, section headings, the 3-column skills, and tab-aligned
    split lines (company/title/project/university). Measured against LibreOffice layout
    (a close proxy for Word). If soffice/pdfplumber are unavailable it does nothing.
    Never changes text, so the integrity gate is unaffected.
    """
    import subprocess, tempfile
    docx_path = Path(docx_path)
    with tempfile.TemporaryDirectory() as td:
        try:
            subprocess.run(["soffice", "--headless", f"-env:UserInstallation=file://{td}/p",
                            "--convert-to", "pdf", "--outdir", td, str(docx_path)],
                           check=True, capture_output=True, timeout=180)
        except Exception:
            return {"applied": [], "note": "soffice unavailable — tightener skipped"}
        pdf = Path(td) / (docx_path.stem + ".pdf")
        if not pdf.exists():
            return {"applied": [], "note": "PDF not produced — tightener skipped"}
        try:
            lines = _pdf_lines(pdf)
        except Exception as e:
            return {"applied": [], "note": f"pdfplumber unavailable — tightener skipped ({e})"}

    from docx import Document
    doc = Document(str(docx_path))
    allp = [p for p in doc.paragraphs if _text(p).strip()]
    page_right = max((x1 for _, _, x1 in lines), default=0.0)

    HEADINGS = ("AREAS OF EXPERTISE", "KEY ACHIEVEMENTS", "PROFESSIONAL EXPERIENCE",
                "PROJECTS", "EDUCATION", "ADDITIONAL SKILLS")

    def _has_tab(p):
        return any(r._r.find(qn("w:tab")) is not None for r in p.runs)

    # paragraphs to skip from tightening (but still used to keep the search moving)
    first3 = {id(p._p) for p in allp[:3]}                       # name / role / contact
    i_ae = next((k for k, p in enumerate(allp) if _text(p).strip().upper().startswith("AREAS OF EXPERTISE")), None)
    i_ka = next((k for k, p in enumerate(allp) if _text(p).strip().upper().startswith("KEY ACHIEVEMENTS")), None)
    skills_ids = set()
    if i_ae is not None and i_ka is not None:
        skills_ids = {id(p._p) for p in allp[i_ae + 1:i_ka]}

    def _is_target(p):
        t = _text(p).strip()
        if id(p._p) in first3 or id(p._p) in skills_ids:
            return False
        if any(t.upper().startswith(h) for h in HEADINGS):
            return False
        if _has_tab(p):
            return False
        return True

    def _locate(pnorm, lb):
        """Find the run of consecutive lines (from index lb) that the paragraph
        occupies. Returns (start, count, widths, x0s) or None."""
        for start in range(lb, len(lines)):
            ln = _norm(lines[start][0])
            if not ln:
                continue
            head = pnorm[:12]
            if not (ln.startswith(head) or pnorm.startswith(ln[:12])):
                continue
            acc, widths, x0s, k = "", [], [], start
            while k < len(lines) and len(acc) < len(pnorm):
                acc += _norm(lines[k][0])
                widths.append(lines[k][2] - lines[k][1])
                x0s.append(lines[k][1])
                k += 1
            if acc.startswith(pnorm) or (pnorm.startswith(acc) and len(acc) >= 0.9 * len(pnorm)):
                return start, k - start, widths, x0s
        return None

    applied = []
    lb = 0
    for p in allp:
        pnorm = _norm(_text(p))
        if not pnorm:
            continue
        loc = _locate(pnorm, lb)
        if loc is None:
            continue                       # can't place it; leave pointer, move on
        start, count, widths, x0s = loc
        lb = start + count                 # only ever advance
        if not _is_target(p) or count != 2:
            continue
        # A bullet's hanging indent (glyph + gap) does NOT scale when the text
        # condenses, so only the TEXT portion is recoverable. The continuation
        # (2nd) visual line starts at the true text margin (no bullet glyph), so
        # x0s[1] gives the text-left; the gap to x0s[0] is the fixed indent.
        text_left = x0s[1] if len(x0s) > 1 else x0s[0]
        indent = max(0.0, text_left - x0s[0])          # fixed, non-scaling
        margin = max(10.0, 0.03 * (page_right - text_left))   # layout-slop safety (Word vs LO)
        avail_text = (page_right - text_left) - margin        # scalable room on one line
        total_text = (widths[0] - indent) + widths[1]         # text-only across both lines
        if avail_text <= 0 or total_text <= 0:
            continue
        s = avail_text / total_text
        # WIDOW_MIN_SCALE is a hard GENTLE cap: never condense harder than it, even
        # on lean tiers whose min_scale would allow more. (max, not min — a higher
        # value is the gentler/less-aggressive bound.) Harder cases wrap instead.
        floor = max(min_scale, WIDOW_MIN_SCALE)
        if s < floor:
            continue                       # genuine multi-line content — leave it
        # This line really wrapped (count == 2). Even when the geometry computes as a
        # bare fit (s >= 1) — the measurement runs a few percent optimistic — nudge it
        # under by SLOP so the trailing widow word pulls up. A couple percent of
        # horizontal scale is visually imperceptible; the floor bounds how far it goes.
        pct = max(int(floor * 100), min(98, int(min(s, 1.0) * 100) - SLOP))
        for r in p.runs:
            _set_char_scale(r, pct)
        applied.append({"scale": pct, "text": _text(p).strip()[:60]})

    if applied:
        doc.save(str(docx_path))
    return {"applied": applied, "min_scale": min_scale,
            "note": f"condensed {len(applied)} near-overflow line(s) to fit one line"}


# --------------------------------------------------------------------------
# Dynamic content budget for the DESIGNED template — PACK THE PAGE.
#
# The template's fonts/margins are fixed (we never resize the design), so the
# only honest lever on length is how much CONTENT we show. This walks a ladder
# of budgets from RICHEST (the most experiences/projects/bullets) down to
# leanest, re-rendering and re-measuring after each, and stops at the FIRST tier
# that fits `max_pages`. Because the ladder is ordered richest-first, "first tier
# that fits" is the *fullest* page the supplied content can produce within the
# page budget — so a resume with lots of relevant material packs the page instead
# of stopping at a lean tier with a half-empty bottom (the white-space problem).
#
# This is the single-pass "grow to fill" design: Filter 3 supplies a GENEROUS
# superset (all genuinely-relevant experiences/projects/bullets, most-relevant
# -first), and the cascade decides how many of them fit. The richer tiers at the
# top of the ladder (rich-max ... full-plus) are what let it grow INTO that
# superset to fill the page; the leaner tiers below are the shrink path for when
# even a tight render still overflows. No second model pass is needed — the
# backfill material is already in `filter3.json`, just held in reserve.
#
# Every tier trims by WHOLE ITEMS only (cap experiences/projects/bullets/areas,
# drop a context blurb / the additional-skills line / a trailing certification)
# and pushes the cosmetic one-line tightener a little harder — it NEVER edits a
# bullet's words or a metric, so integrity is untouched (omission, never
# fabrication). Items are most-relevant-first, so trimming drops the least
# -relevant material and growing adds the next-most-relevant.
#
# The winning render's page-fill ratio is measured and reported (`fill_ratio`,
# `whitespace_in`) so an under-filled page is visible — that means the CV simply
# ran out of relevant material at this density, not that the cascade quit early.
#
# If even the tightest tier is still over target, the tightest render is left on
# disk and fit_ok=False is reported so a human decides (accept the overflow, or
# cut more upstream in Filter 3). If page count can't be measured (no
# LibreOffice), it renders the full budget once and reports fit_ok=None.
# --------------------------------------------------------------------------

from dataclasses import dataclass as _dataclass, replace as _replace  # noqa: E402


@_dataclass(frozen=True)
class BudgetTier:
    name: str
    key_achievements: int   # max KEY ACHIEVEMENTS kept
    max_experiences: int    # max roles kept
    top_bullets: int        # max bullets on the most-relevant role
    other_bullets: int      # max bullets on every other role
    max_projects: int       # max projects kept
    min_scale: float        # tightener floor (lower = condense harder)
    # --- header (presentation-overhead) levers, trimmed BEFORE content floors ---
    areas_cap: int = 99     # max AREAS OF EXPERTISE items kept (3-col block height)
    drop_context: bool = False  # drop experiences' one-line `context` blurb
    drop_additional: bool = False  # drop the ADDITIONAL SKILLS line + heading
    max_certifications: int = 99    # max certifications kept (trailing, least-relevant dropped first)


# Richest -> leanest. The cascade returns the FIRST (richest) tier that fits the
# page budget, so the top of the ladder is what packs a content-rich CV onto a
# full page. The CONTENT FLOORS (>=3 experiences, >=2 projects, top role >=3
# bullets) are protected for as long as possible: every tier from `rich-max`
# down through `lean-max` honors all three floors, trimming only richness above
# the floor (extra experiences/projects/bullets) and then HEADER overhead
# (areas-of-expertise count, context lines, additional-skills line, trailing
# certs). Only the final three tiers drop BELOW a floor (one project, then 2 top
# bullets, then 2 experiences) — a true last resort when even a fully-lean header
# still overflows. Header is cut before content; richness is grown before header.
#
# `rich-*` and `full-plus` exist to GROW into a generous Filter-3 superset and
# fill the page; `apply_budget_tier` caps by slicing, so a tier asking for more
# items than the CV supplies simply yields what's there (graceful no-op).
BUDGET_LADDER = [
    BudgetTier("rich-max",     key_achievements=3, max_experiences=5, top_bullets=6, other_bullets=3, max_projects=4, min_scale=0.94, areas_cap=9, drop_context=False, drop_additional=False, max_certifications=99),
    BudgetTier("rich",         key_achievements=3, max_experiences=5, top_bullets=5, other_bullets=3, max_projects=3, min_scale=0.93, areas_cap=9, drop_context=False, drop_additional=False, max_certifications=99),
    BudgetTier("rich-lean",    key_achievements=3, max_experiences=4, top_bullets=5, other_bullets=3, max_projects=3, min_scale=0.93, areas_cap=9, drop_context=False, drop_additional=False, max_certifications=99),
    BudgetTier("full-plus",    key_achievements=3, max_experiences=4, top_bullets=5, other_bullets=2, max_projects=2, min_scale=0.92, areas_cap=9, drop_context=False, drop_additional=False, max_certifications=99),
    BudgetTier("full",         key_achievements=3, max_experiences=3, top_bullets=4, other_bullets=2, max_projects=2, min_scale=0.92, areas_cap=9, drop_context=False, drop_additional=False, max_certifications=99),
    BudgetTier("lean-context", key_achievements=3, max_experiences=3, top_bullets=4, other_bullets=2, max_projects=2, min_scale=0.90, areas_cap=8, drop_context=True,  drop_additional=False, max_certifications=99),
    BudgetTier("lean-ka",      key_achievements=2, max_experiences=3, top_bullets=4, other_bullets=2, max_projects=2, min_scale=0.90, areas_cap=6, drop_context=True,  drop_additional=False, max_certifications=2),
    BudgetTier("lean-areas",   key_achievements=2, max_experiences=3, top_bullets=3, other_bullets=2, max_projects=2, min_scale=0.88, areas_cap=5, drop_context=True,  drop_additional=True,  max_certifications=1),
    BudgetTier("lean-bullets", key_achievements=2, max_experiences=3, top_bullets=3, other_bullets=1, max_projects=2, min_scale=0.86, areas_cap=4, drop_context=True,  drop_additional=True,  max_certifications=1),
    BudgetTier("lean-max",     key_achievements=2, max_experiences=3, top_bullets=3, other_bullets=1, max_projects=2, min_scale=0.84, areas_cap=3, drop_context=True,  drop_additional=True,  max_certifications=1),
    # --- below here, content floors are broken only because a lean header still overflows ---
    BudgetTier("one-project",  key_achievements=2, max_experiences=3, top_bullets=3, other_bullets=1, max_projects=1, min_scale=0.84, areas_cap=3, drop_context=True,  drop_additional=True,  max_certifications=1),
    BudgetTier("hard-bullets", key_achievements=2, max_experiences=3, top_bullets=2, other_bullets=1, max_projects=1, min_scale=0.82, areas_cap=3, drop_context=True,  drop_additional=True,  max_certifications=1),
    BudgetTier("last-resort",  key_achievements=2, max_experiences=2, top_bullets=2, other_bullets=1, max_projects=1, min_scale=0.80, areas_cap=3, drop_context=True,  drop_additional=True,  max_certifications=1),
]


# Deployment policy: certifications that must appear on EVERY resume regardless of
# role or how lean the page budget gets. Matched case-insensitively as substrings
# against each cert's text. Pinned certs are floated to the FRONT and are EXEMPT from
# the per-tier certification cap, so the cascade can never trim them away. (Honest:
# these are real cv.md credentials the user has deemed universally relevant. The
# filter must still PUT the cert in the JSON — this only guarantees the cap can't drop
# it.) Set to () to disable pinning.
PINNED_CERT_PATTERNS = ("pmp", "project management professional")


def _is_pinned_cert(cert: str) -> bool:
    c = (cert or "").lower()
    return any(pat in c for pat in PINNED_CERT_PATTERNS)


def apply_budget_tier(tailored: dict, tier: BudgetTier) -> dict:
    """Return a deep copy of `tailored` trimmed to `tier` by whole-item omission.
    Caps counts only — never edits any text or metric.

    Header overhead is trimmed too (areas-of-expertise count, experience context
    lines), since the cascade prefers shedding presentation overhead before it
    drops below the content floors. These are still whole-item omissions: a
    trailing area is dropped, a context blurb is removed in full — no wording is
    rewritten, so the integrity gate result is unchanged."""
    d = copy.deepcopy(tailored)
    if isinstance(d.get("key_achievements"), list):
        d["key_achievements"] = d["key_achievements"][:tier.key_achievements]
    exps = (d.get("experience") or [])[:tier.max_experiences]
    for i, role in enumerate(exps):
        cap = tier.top_bullets if i == 0 else tier.other_bullets
        if isinstance(role.get("bullets"), list):
            role["bullets"] = role["bullets"][:cap]
        if tier.drop_context and "context" in role:
            role["context"] = ""
    d["experience"] = exps
    if isinstance(d.get("projects"), list):
        d["projects"] = d["projects"][:tier.max_projects]
    if isinstance(d.get("areas_of_expertise"), list):
        d["areas_of_expertise"] = d["areas_of_expertise"][:tier.areas_cap]
    if isinstance(d.get("certifications"), list):
        # Pinned certs (deployment policy) always survive and lead; the per-tier cap
        # applies only to the remainder. Whole-item selection — no text is edited.
        certs = d["certifications"]
        pinned = [c for c in certs if _is_pinned_cert(c)]
        rest = [c for c in certs if not _is_pinned_cert(c)]
        keep_rest = rest[:max(0, tier.max_certifications - len(pinned))]
        d["certifications"] = pinned + keep_rest
    if tier.drop_additional:
        # empty BOTH the explicit line and the derive source so fill() renders nothing
        d["additional_skills"] = ""
        d["skills"] = []
    return d


def _TIER(name: str) -> "BudgetTier":
    """Look up a ladder tier by name (used for sane fallbacks)."""
    return next(b for b in BUDGET_LADDER if b.name == name)


def _grow_ops():
    """Priority order for Phase-2 backfill, as (d_exp, d_top_bullet, d_other_bullet,
    d_project) unit adds. EXPERIENCES and PROJECTS first (the user's stated
    preference for what should fill the page), then bullets. The fit loop stays on
    each op while it keeps fitting, so a single entry adds as many of that item as
    the page allows before moving on."""
    return [(1, 0, 0, 0),   # + one experience
            (0, 0, 0, 1),   # + one project
            (0, 1, 0, 0),   # + one bullet on the top role
            (0, 0, 1, 0)]   # + one bullet on each other role


_GROW_CAPS = (8, 8, 6, 8)   # ceilings for (max_experiences, top_bullets, other_bullets, max_projects)


def render_fit_template(tailored: dict, out_path: Path, template_path: Path,
                        max_pages: int = 1, no_tighten: bool = False,
                        no_grow: bool = False, fill_target: float = 0.96) -> dict:
    """Two-phase 'trim to fit, then fill' render.

    PHASE 1 (shrink): walk BUDGET_LADDER richest-first, trimming header overhead
    (then, as a last resort, content floors) until the resume fits `max_pages`.
    PHASE 2 (grow): holding that winning tier's HEADER settings fixed, grow the
    CONTENT back from the Filter-3 superset by GREEDY ACCRETION — add a whole
    experience, then a project, then bullets (experiences/projects first, the
    user's preference), keeping each add that still fits. This is what fills the
    bottom white space: Phase 1 finds the header budget, Phase 2 packs the page
    with as much real content as fits under it.

    Phase 2 never re-inflates the header (so the 'header is cut before content'
    guarantee holds), never runs once Phase 1 had to break a content floor (we're
    already short on space), and skips when Phase 1's page is already
    >= `fill_target` full. Growth is whole-item only (caps), so integrity is
    untouched. Disable Phase 2 with `no_grow`."""
    # Page-fit measurement runs through measure_fit, which uses real MS Word when
    # available (authoritative — matches what the candidate sees) and falls back to
    # LibreOffice otherwise. So the whole shrink/grow cascade targets the engine the
    # resume is actually viewed in.
    out_path = Path(out_path)
    base = {"key_achievements": len(tailored.get("key_achievements") or []),
            "experiences": len(tailored.get("experience") or []),
            "projects": len(tailored.get("projects") or [])}
    attempts: list[dict] = []

    def _render(tier):
        """Fill + tighten + measure `tier` onto out_path. Returns (status, pages, trimmed)."""
        trimmed = apply_budget_tier(tailored, tier)
        status = fill(trimmed, out_path, template_path)
        if not no_tighten:
            status["tighten"] = tighten_wrapped_lines(out_path, min_scale=tier.min_scale)
        return status, measure_fit.pages(out_path), trimmed

    def _trim_note(trimmed):
        kept = {"key_achievements": len(trimmed.get("key_achievements") or []),
                "experiences": len(trimmed.get("experience") or []),
                "projects": len(trimmed.get("projects") or [])}
        return {k: f"{base[k]}->{kept[k]}" for k in base if kept[k] != base[k]}

    # ----- PHASE 1: shrink header (then floors) until it fits -----
    win = None  # (tier, status, pages, trimmed)
    for tier in BUDGET_LADDER:
        status, pages, trimmed = _render(tier)
        if pages is None:
            # Can't measure -> don't trim blindly. Render the FULL budget and flag.
            full = fill(apply_budget_tier(tailored, _TIER("full")), out_path, template_path)
            if not no_tighten:
                full["tighten"] = tighten_wrapped_lines(out_path, min_scale=_TIER("full").min_scale)
            full.update({"pages": None, "fit_ok": None, "tier": "full", "max_pages": max_pages,
                         "note": "page count unmeasurable (need LibreOffice + pypdf/pdfinfo); "
                                 "rendered at the full budget without verifying length"})
            return full
        attempts.append({"tier": tier.name, "pages": pages, "trimmed": _trim_note(trimmed) or "none"})
        if pages <= max_pages:
            win = (tier, status, pages, trimmed)
            break

    if win is None:
        # Even the tightest tier overflows. Leave the tightest render on disk + flag.
        status, pages, trimmed = _render(BUDGET_LADDER[-1])
        status.update({
            "pages": pages, "fit_ok": False, "tier": BUDGET_LADDER[-1].name,
            "max_pages": max_pages, "ladder": attempts,
            "note": (f"still {pages} page(s) at the tightest budget "
                     f"('{BUDGET_LADDER[-1].name}'). Do NOT shrink the template — accept the "
                     f"overflow and mark 'needs human check', or cut content upstream in Filter 3."),
        })
        fi = measure_fit.fill(out_path)
        if fi:
            status["fill"] = fi
        return status

    win_tier, win_status, win_pages, win_trimmed = win
    win_fill = measure_fit.fill(out_path)        # baseline (Phase 1) render is on disk

    # ----- PHASE 2: greedily grow content under the winning header to fill -----
    floors_ok = (win_tier.max_experiences >= 3 and win_tier.max_projects >= 2
                 and win_tier.top_bullets >= 3)
    under_filled = (not win_fill) or win_fill.get("fill_ratio", 1.0) < fill_target
    if (not no_grow) and floors_ok and under_filled:
        base_tier = win_tier
        base_name = win_tier.name
        base_content = [base_tier.max_experiences, base_tier.top_bullets,
                        base_tier.other_bullets, base_tier.max_projects]
        cur = list(base_content)
        good = list(base_content)        # content of the last FITTING render
        ops = _grow_ops()
        budget = 16                      # bound the render count
        i = 0
        while i < len(ops) and budget > 0:
            d = ops[i]
            cand = [min(cur[k] + d[k], _GROW_CAPS[k]) for k in range(4)]
            if cand == cur:              # already at supply/ceiling for this op
                i += 1
                continue
            gt = _replace(base_tier, name=base_name + "+fill",
                          max_experiences=cand[0], top_bullets=cand[1],
                          other_bullets=cand[2], max_projects=cand[3])
            _, g_pages, g_trimmed = _render(gt)
            budget -= 1
            attempts.append({"tier": gt.name, "pages": g_pages,
                             "trimmed": _trim_note(g_trimmed) or "none"})
            if g_pages is not None and g_pages <= max_pages:
                cur = cand
                good = list(cur)         # accept; stay on this op to add more while it fits
            else:
                i += 1                   # this add overflows; try a smaller-impact op
        # COMMIT: re-render the best fitting content once, so the on-disk file is
        # correct (the search may have ended on an overflow) and the tier name is clean.
        grew = good != base_content
        commit = _replace(base_tier, name=(base_name + "+fill") if grew else base_name,
                          max_experiences=good[0], top_bullets=good[1],
                          other_bullets=good[2], max_projects=good[3])
        win_status, win_pages, win_trimmed = _render(commit)
        win_tier = commit
        win_fill = measure_fit.fill(out_path)

    win_status.update({
        "pages": win_pages, "fit_ok": True, "tier": win_tier.name,
        "max_pages": max_pages, "trimmed": _trim_note(win_trimmed) or "none",
        "ladder": attempts,
    })
    if win_fill:
        win_status["fill"] = win_fill
        if win_fill["fill_ratio"] < 0.9:
            win_status["note"] = (
                f"page is {int(win_fill['fill_ratio']*100)}% full at tier '{win_tier.name}'; "
                f"~{win_fill['whitespace_in']}in of bottom white space remains because the "
                f"supplied content doesn't fill the page even after growth — add more relevant "
                f"experiences/projects/bullets to filter3.json if you want it fuller.")
    return win_status


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Fill template_0.docx with a TailoredResume")
    ap.add_argument("--tailored", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--template", default=str(common.TEMPLATE_PATH))
    ap.add_argument("--no-tighten", action="store_true",
                    help="skip the one-line tightener (condensing near-overflow lines)")
    ap.add_argument("--min-scale", type=float, default=0.90,
                    help="tightener floor when NOT cascading (with --no-cascade); "
                         "lines needing more condense than this are left to wrap")
    ap.add_argument("--max-pages", type=int, default=1,
                    help="page target for the dynamic content budget (default 1). The "
                         "renderer trims by whole items across the budget ladder until it fits.")
    ap.add_argument("--no-cascade", action="store_true",
                    help="disable the dynamic budget; fill once at the full budget + tighten")
    ap.add_argument("--no-grow", action="store_true",
                    help="disable Phase 2 (content backfill); just shrink-to-fit")
    ap.add_argument("--fill-target", type=float, default=0.96,
                    help="Phase 2 stops growing once the page is this full (0..1, default 0.96)")
    args = ap.parse_args(argv)
    tailored = common.load_json(args.tailored)
    if args.no_cascade:
        status = fill(tailored, Path(args.out), Path(args.template))
        if not args.no_tighten:
            status["tighten"] = tighten_wrapped_lines(Path(args.out), min_scale=args.min_scale)
    else:
        status = render_fit_template(tailored, Path(args.out), Path(args.template),
                                     max_pages=args.max_pages, no_tighten=args.no_tighten,
                                     no_grow=args.no_grow, fill_target=args.fill_target)
    print(json.dumps(status, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
